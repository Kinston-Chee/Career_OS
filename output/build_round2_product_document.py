from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

OUT = Path("outputs/documentation/CareerOS_Product_Documentation_Round_2.docx")
OUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = "102A43"; BLUE = "2E74B5"; CYAN = "4CB9D8"; PURPLE = "6D5BD0"
LIGHT = "EEF4FA"; LAV = "F1EEFF"; GRAY = "5F6B7A"; BORDER = "D7E1EC"; WHITE = "FFFFFF"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(.49)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = RGBColor.from_string("243447")
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
for name, size, before, after, color in [
    ("Title", 30, 0, 8, NAVY), ("Subtitle", 14, 0, 10, GRAY),
    ("Heading 1", 17, 16, 8, BLUE), ("Heading 2", 13, 12, 6, BLUE),
    ("Heading 3", 11.5, 8, 4, NAVY)
]:
    s = styles[name]; s.font.name = "Calibri"; s.font.size = Pt(size); s.font.color.rgb = RGBColor.from_string(color)
    s.font.bold = name != "Subtitle"; s.paragraph_format.space_before = Pt(before); s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

if "Callout" not in styles:
    s = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    s.font.name = "Calibri"; s.font.size = Pt(11); s.font.bold = True; s.font.color.rgb = RGBColor.from_string(NAVY)
    s.paragraph_format.space_before = Pt(6); s.paragraph_format.space_after = Pt(8)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd")) or OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    if shd.getparent() is None: tcPr.append(shd)

def margins(cell, top=80, start=120, bottom=80, end=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar") or OxmlElement("w:tcMar")
    if tcMar.getparent() is None: tcPr.append(tcMar)
    for tag, val in (("top",top),("start",start),("bottom",bottom),("end",end)):
        el = tcMar.find(qn("w:"+tag)) or OxmlElement("w:"+tag)
        el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
        if el.getparent() is None: tcMar.append(el)

def set_cell_text(cell, text, bold=False, color=None, size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); r.bold = bold; r.font.name = "Calibri"; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    trPr = t.rows[0]._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader"); tblHeader.set(qn("w:val"), "true"); trPr.append(tblHeader)
    for i,h in enumerate(headers):
        shade(t.rows[0].cells[i], NAVY); margins(t.rows[0].cells[i])
        set_cell_text(t.rows[0].cells[i], h, True, WHITE, 9)
    for ri,row in enumerate(rows):
        new_row = t.add_row()
        trPr = new_row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit"); trPr.append(cantSplit)
        cells = new_row.cells
        for i,v in enumerate(row):
            if ri % 2: shade(cells[i], "F8FAFC")
            margins(cells[i]); set_cell_text(cells[i], str(v), False, None, 9.2)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(.5 + level*.25)
    p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(.5); p.paragraph_format.first_line_indent = Inches(-.25)
    p.paragraph_format.space_after = Pt(4); p.add_run(text); return p

def callout(label, text, fill=LIGHT):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = False
    c=t.cell(0,0); c.width=Inches(6.5); shade(c,fill); margins(c,140,180,140,180)
    c.text=""; p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    r=p.add_run(label+"  "); r.bold=True; r.font.color.rgb=RGBColor.from_string(BLUE)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def h1(text): doc.add_heading(text,1)
def h2(text): doc.add_heading(text,2)
def h3(text): doc.add_heading(text,3)
def para(text, bold_lead=None):
    p=doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold=True; p.add_run(text[len(bold_lead):])
    else: p.add_run(text)
    return p
def page(): doc.add_page_break()

# Running furniture
header = sec.header.paragraphs[0]
header.text = "CAREEROS  ·  ROUND 2 PRODUCT DOCUMENTATION"
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.runs[0].font.size = Pt(8); header.runs[0].font.bold=True; header.runs[0].font.color.rgb=RGBColor.from_string(GRAY)
footer = sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Talentbank AI Hackathon 2026  ·  Product documentation").font.size=Pt(8)

# Cover
doc.add_paragraph().paragraph_format.space_after = Pt(110)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("CAREEROS"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(PURPLE); r.font.letter_spacing if False else None
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("CareerOS"); r.bold=True; r.font.size=Pt(34); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The evidence memory and action layer\nbetween education and employment"); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
r=p.add_run("ROUND 2 PRODUCT DOCUMENTATION"); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(PURPLE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Talentbank AI Hackathon 2026  ·  July 2026"); r.font.size=Pt(10); r.font.color.rgb=RGBColor.from_string(GRAY)
doc.add_paragraph().paragraph_format.space_after=Pt(115)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Remember the signal. Prove it. Act earlier."); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string(NAVY)
page()

h1("Document Purpose and Reading Guide")
para("This document replaces the Phase 1 product description with a Round 2 view grounded in the current CareerOS repository, implemented routes, local state flows, partial backend, and the deeper product-value analysis completed for the submission.")
callout("Truth boundary", "CareerOS is currently a polished multi-workspace prototype. It demonstrates product logic and rich interactions using local or mock data, deterministic recommendations, optional model calls, and a partial FastAPI backend. It does not yet prove adoption, customer demand, AI accuracy, production integrations, autonomous agents, or measured impact.", LAV)
h2("Contents")
for x in [
"1. Executive Summary","2. Product Thesis and Problem","3. Users, Buyer, and Pilot Wedge",
"4. CareerOS Operating Model","5. Student Workspace","6. Employer Workspace","7. University Workspace",
"8. Cross-Workspace Evidence Loop","9. Implementation Status","10. Technical Architecture",
"11. AI, Data, and Governance","12. Strategic Value for Talentbank","13. Pilot and Validation Plan",
"14. Product Roadmap","15. Setup and Demo Guidance","Appendix A. Route and Capability Inventory",
"Appendix B. Claim Language"
]: bullet(x)
page()

h1("1. Executive Summary")
para("Career potential is frequently lost between moments. A student completes a project but cannot later explain what it proves. An employer meets strong early talent but restarts sourcing when an urgent role appears. A university identifies a capability gap after the current cohort has already lost time.")
para("CareerOS is designed to preserve those signals as evidence and turn them into earlier decisions and actions. It connects three role-specific workspaces—Student, Employer, and University—through a common operating logic: signal → evidence → decision → action → outcome → learning.")
callout("Positioning", "CareerOS is the evidence memory and action layer between education and employment. It is more than a job portal, chatbot, dashboard, or isolated matching tool because its central value is continuity across decisions.")
h2("What changed in Round 2")
for x in [
"The student experience expanded beyond profiles and listings into Career Memory, Career Intelligence, skill development, opportunities, applications, interview practice, communities, mentorships, and an AI Companion.",
"The employer experience expanded into an AI-first home, explainable talent discovery, candidate review, warm campus pipeline, engagement creation, hiring action queues, marketplace workflows, and analytics.",
"The university experience expanded into Student Readiness, Curriculum–Market Alignment, Alumni Signals, Collaboration Marketplace, Accreditation Hub, and an AI Office with human decision controls.",
"Cross-page stores now demonstrate selected evidence handoffs, shortlist and pipeline state, curriculum evidence packs, interventions, and accreditation readiness changes.",
"The documentation now separates demonstrated, simulated/partial, and future-facing capabilities."
]: bullet(x)

h1("2. Product Thesis and Problem")
h2("2.1 One structural failure, three stakeholder consequences")
table(["Stakeholder","Failure today","Consequence"],[
("Student","Experience, feedback, and proof are fragmented.","Uncertainty, mass applications, random learning, and loss of confidence."),
("Employer","Early-talent interactions and decision context decay across events, spreadsheets, ATSs, and inboxes.","Repeated sourcing, weak validation, and lost warm candidates."),
("University","Readiness, market, alumni, partnership, and accreditation evidence arrives late or remains siloed.","Intervention after the current cohort can benefit and repeated evidence assembly.")
],[1.2,2.6,2.7])
h2("2.2 Existing tools optimise transactions")
para("Job boards optimise vacancies, ATS products optimise applicant processing, LMS products optimise course delivery, and reporting systems describe outcomes. These tools may remain important. CareerOS is designed to preserve the context between them so the next decision does not begin from zero.")
h2("2.3 Founder-level thesis")
callout("Core thesis", "CareerOS protects career potential from being lost between moments. Career Memory prevents evidence loss. Career Intelligence prevents option loss. Employer reactivation prevents relationship loss. University interventions prevent time loss. Accreditation reuse prevents institutional knowledge loss.")

h1("3. Users, Buyer, and Pilot Wedge")
h2("3.1 Primary users")
table(["Workspace","Representative user","Primary job"],[
("Student","Chris Lee, Year 3 Data Science student at Taylor’s University","Preserve evidence, explore paths, build gaps, prepare, apply, and recover."),
("Employer","Recruiter, early-talent lead, campus-relations or hiring manager","Discover evidence-backed talent, validate fit, preserve relationships, and act on hiring priorities."),
("University","Faculty leader, programme director, career-services or employability owner","Detect readiness and market gaps, assign interventions, coordinate partners, and reuse evidence.")
],[1.1,2.2,3.2])
h2("3.2 Recommended first buyer and wedge")
para("The strongest initial buyer hypothesis is a faculty-level employability owner, sponsored by university leadership. The first product should not attempt to launch the entire three-sided network at once.")
callout("Recommended pilot", "One computing or data faculty, one cohort, two priority capability gaps, one career-services owner, and one or two employer-backed interventions over 8–12 weeks.")
h2("3.3 Why this wedge")
for x in [
"The university owns an urgent operational problem and can coordinate a bounded cohort.",
"The workflow creates direct value before cross-organisation network effects exist.",
"The intervention naturally activates student evidence and employer participation.",
"Manual, consented data import is sufficient for an initial validation; full SIS, LMS, ATS, or HRIS integration is not required."
]: bullet(x)

h1("4. CareerOS Operating Model")
h2("4.1 Evidence-to-action sequence")
table(["Stage","CareerOS meaning","Examples"],[
("Signal","A moment that may matter later.","Project, challenge, application, hiring decision, readiness gap, alumni outcome."),
("Evidence","The signal with source, context, recency, and supporting material.","Career Memory entry, challenge observation, curriculum evidence, intervention record."),
("Decision","A prioritised judgement that exposes uncertainty.","Next skill, validation question, shortlist, intervention, programme priority."),
("Action","A specific follow-through state.","Practise, apply, re-engage, assign, collaborate, approve."),
("Outcome","The observable result of the action.","Completion, progression, candidate movement, intervention result."),
("Learning","Context preserved for the next decision.","Updated guidance, warmer pipeline, programme evidence, accreditation reuse.")
],[.85,2.55,3.1])
h2("4.2 Three workspaces, one product logic")
para("Each workspace addresses a different expression of evidence loss, but they are not intended to be independent dashboards. Student evidence can support employer validation; employer interactions can create new student evidence; institutional gaps can trigger interventions that create both candidate signals and reusable university records.")
h2("4.3 What CareerGraph means in Round 2")
para("CareerGraph is the product concept for connecting skills, experiences, roles, opportunities, employer demand, interventions, and outcomes. In the current prototype, this logic is represented through route-level interactions, local stores, mock datasets, and selected API calls. It is not yet a production shared graph database or live cross-organisation intelligence network.")

h1("5. Student Workspace")
para("The Round 2 student experience is centred on a Data Science student and a continuous evidence-to-action journey rather than a one-time job search.")
h2("5.1 Career Memory")
para("Career Memory preserves projects, internships, leadership, events, achievements, reflections, skills, and supporting evidence while the context is still available. Students can add or improve entries and see how experiences contribute to an evolving career story.")
callout("Designed behaviour change", "Maintain reusable evidence continuously instead of reconstructing a CV only when an application appears.")
h2("5.2 Career Intelligence")
para("Career Intelligence interprets existing evidence against possible paths, including adjacent roles beyond the degree title. It helps a student identify transferable strengths, missing proof, and a prioritised next step.")
h2("5.3 Skill development, opportunities, and applications")
for x in [
"Skill Development links an identified gap to a roadmap, learning activity, project, or evidence-building task.",
"Opportunities explain why an event or role may fit and support saving, preparation, and application flows.",
"Applications preserve stage, deadlines, and the next action across an application pipeline."
]: bullet(x)
h2("5.4 AI Companion and interview practice")
para("The AI Companion provides intent-based guidance and routes the student to relevant context. Interview Practice supports role- and company-specific preparation, including voice-session interaction patterns. Optional backend/model paths exist, but scripted fallbacks and simulated content remain important parts of the demo.")
h2("5.5 Community, mentorship, and lifecycle return")
para("Communities, mentorships, notifications, and recommendations create reasons to return between applications. The community experience is comparatively rich; mentorship and a complete rejection-to-recovery loop remain less mature and should be described accordingly.")

h1("6. Employer Workspace")
h2("6.1 Employer Home")
para("The employer home acts as an AI briefing and action surface. It summarises operational signals, candidates, pipeline status, opportunities, and next actions without changing the underlying employer workflow.")
h2("6.2 Explainable talent discovery")
para("Talent Discovery starts from a hiring need and presents candidate context, supporting evidence, strengths, gaps, and targeted validation questions. The interface models evidence-aware selection rather than opaque ranking.")
callout("Truth boundary", "Candidate records, match percentages, confidence indicators, and evidence histories are demo data. Production provenance, fairness evaluation, and validated matching accuracy are not implemented.")
h2("6.3 Candidate review and decision history")
para("Candidate pages preserve shortlisting, pass decisions, interview preparation questions, and rationale through local state. The intended behaviour is to make decisions explainable and challengeable, rather than to treat a score as proof.")
h2("6.4 Campus pipeline and warm reactivation")
para("Campus Pipeline demonstrates how prior challenges, events, workshops, or hiring processes could become a reusable warm pool. When an urgent need appears, employers can revisit known context instead of restarting sourcing from zero.")
h2("6.5 Engagements, marketplace, and analytics")
for x in [
"Engagement creation supports structured challenges, workshops, and campus activities designed to reveal applied ability.",
"The marketplace and applicant workflows model role creation, applicant stages, and candidate movement.",
"Analytics and briefing views show how employer activity could create reusable operational learning; displayed outcomes remain illustrative."
]: bullet(x)

h1("7. University Workspace")
h2("7.1 AI-first university overview")
para("The Round 2 overview frames the university as an employability operating system. It surfaces programme, cohort, market, partnership, accreditation, and intervention signals through action-oriented cards and natural-language routing.")
h2("7.2 Student Readiness")
para("Student Readiness highlights hidden employability risk, differentiates academic performance from career readiness, and supports assignment of intervention owners, deadlines, and follow-up states.")
h2("7.3 Curriculum–Market Alignment")
para("Curriculum–Market Alignment compares programme coverage with illustrative demand signals, provides evidence chains, and supports current-cohort interventions while formal curriculum changes continue. Evidence packs and roadmap actions demonstrate how analysis becomes an operational artefact.")
h2("7.4 Alumni Signals and Collaboration Marketplace")
para("Alumni Signals route graduate outcomes and employer feedback back into programme discussions. Collaboration Marketplace connects identified capability gaps to employer, society, workshop, challenge, and outreach workflows.")
h2("7.5 Accreditation Hub")
para("Accreditation Hub models continuous evidence readiness, named ownership, requirement tracking, and reuse across MQA, SETARA, QS, and internal review contexts. It demonstrates preparation and reuse—not automated accreditation submission.")
h2("7.6 University AI Office")
para("The AI Office and Decision Room present specialised assistant roles that prepare cross-department actions. Human approval remains explicit. Some optional model calls are configured in the frontend, but the operating-room intelligence, metrics, and external effects are simulated.")

h1("8. Cross-Workspace Evidence Loop")
h2("8.1 Example: one capability gap")
for x in [
"A university identifies a cloud or MLOps readiness gap affecting the current cohort.",
"The programme team selects a targeted employer-backed challenge or workshop while curriculum reform proceeds.",
"Students participate and demonstrate applied ability.",
"The activity can become Career Memory evidence with source and context.",
"An employer can use the evidence as a discovery and validation signal.",
"The university measures participation and follow-through, then reuses the record in programme and accreditation work."
]: numbered(x)
callout("Round 2 distinction", "The current repository demonstrates many components of this loop and selected state handoffs. The complete cross-organisation evidence network, permissions, and outcome propagation remain future-facing.")
h2("8.2 Why the loop matters")
table(["Stakeholder","Value that can accumulate"],[
("Student","A portable career story and clearer evidence gaps."),
("Employer","Relationship history, observed signals, and reusable candidate context."),
("University","Intervention history, outcome learning, and reusable institutional evidence."),
("Talentbank","A potential lifecycle relationship beyond listings and one-off applications.")
],[1.2,5.3])

h1("9. Implementation Status")
table(["Capability","Current status","Boundary"],[
("Three role-based workspaces","Demonstrated","Protected routes and rich local interactions are implemented."),
("Career Memory and profile state","Demonstrated / partial","Local state and selected extraction paths; production verification is incomplete."),
("Career Intelligence and paths","Demonstrated / simulated","Interactive paths and recommendations use illustrative data."),
("Opportunities and applications","Demonstrated locally","No live job marketplace or production application exchange."),
("Interview practice and AI Companion","Demonstrated / optional AI","Backend or model paths may be used; fallbacks and scripted states remain."),
("Communities and mentorships","Mixed","Community is richer; mentorship lifecycle is thinner."),
("Employer discovery and pipeline","Demonstrated / simulated data","No validated matching, fairness proof, or production ATS integration."),
("University readiness and curriculum workflows","Demonstrated / simulated data","Signals, scores, and impact values are illustrative."),
("Accreditation and AI Office","Demonstrated interaction model","Formal integrations, autonomous action, and production governance are future."),
("Cross-workspace network","Partial concept proof","Selected stores and handoffs exist; live shared intelligence does not.")
],[1.9,1.35,3.25])
h2("9.1 Claims removed or corrected from Round 1")
for x in [
"Removed unsupported market-size and prevalence statistics.",
"Replaced “verified evidence” with evidence-oriented or verification-model language.",
"Removed the claim that Google ADK multi-agent orchestration is fully implemented across the product.",
"Removed production Cloud Run, Cloud SQL, OAuth, SIS/LMS/ATS, and labour-market integration claims.",
"Replaced predictive accuracy and measured-impact language with designed mechanisms and pilot metrics.",
"Reframed competitive uniqueness and business-model claims as hypotheses requiring validation."
]: bullet(x)

h1("10. Technical Architecture")
h2("10.1 Frontend")
table(["Layer","Current implementation"],[
("Framework","React 19 with Vite 5."),
("Routing","React Router with role-protected Student, Employer, and University routes."),
("Styling","Tailwind CSS plus application-specific styles."),
("State","Zustand stores and local component state for selected cross-page workflows."),
("Visualisation","React Force Graph and custom React components."),
("Icons","Lucide React.")
],[1.6,4.9])
h2("10.2 Backend")
para("The repository contains a partial FastAPI and SQLAlchemy backend with candidate, employer, and user routers, database models, migrations, and utilities. Candidate-facing service code references chat, interview-session, and add-experience paths. The frontend also includes local fallbacks when a backend endpoint or environment configuration is unavailable.")
h2("10.3 Intelligence paths")
for x in [
"Deterministic and mock datasets drive many recommendations, scores, summaries, and workflow states.",
"Selected frontend services can call candidate chat, interview, and experience-extraction endpoints.",
"The University AI Office contains an optional direct model configuration path.",
"Backend agent code and sub-agent experiments exist, but they should be treated as partial technical exploration rather than a complete production orchestration layer."
]: bullet(x)
h2("10.4 Production requirements")
para("A production release would require secure identity, role and consent enforcement, durable storage, audit logs, provenance, data minimisation, integration contracts, model evaluation, observability, deployment hardening, and institutional governance.")

h1("11. AI, Data, and Governance")
h2("11.1 AI product principle")
para("AI should reduce navigation, synthesise context, and prepare decisions. It should not conceal uncertainty, substitute an opaque score for evidence, or execute consequential external actions without authority.")
h2("11.2 Human-governed interaction")
table(["AI may prepare","Humans must control"],[
("A next-action recommendation","Whether the student accepts or shares it."),
("Candidate fit rationale and validation questions","Shortlist, rejection, interview, and hiring decisions."),
("Intervention or curriculum options","Programme priorities, owners, budgets, and implementation."),
("Accreditation evidence drafts","Formal claims, approval, and submission."),
("Cross-department action brief","Any external or consequential action.")
],[3.1,3.4])
h2("11.3 Data and evidence requirements")
for x in [
"Student ownership, consent, correction, revocation, and visibility rules for each Career Memory item.",
"Source, provenance, recency, confidence, and verification state for evidence.",
"Clear separation between observed evidence, inferred recommendation, and user-entered claim.",
"Role-based access, minimisation, retention, appeal, and auditability.",
"Fairness and outcome validation before predictive or ranking claims."
]: bullet(x)

h1("12. Strategic Value for Talentbank")
para("CareerOS could extend Talentbank’s relevance from individual career transactions into a longer evidence-and-action lifecycle.")
table(["Dimension","Potential strategic value","What must be validated"],[
("Student lifecycle","Engagement before, during, and after applications.","Retention triggers, consent, and sustained evidence maintenance."),
("Employer offering","Discovery, warm pipeline recovery, engagement design, and reusable talent context.","Buyer urgency, trust, supply, and workflow integration."),
("University offering","A faculty-level employability operating layer.","Budget ownership, implementation capacity, and measurable workflow value."),
("Network learning","Challenges and interventions create new evidence and relationships.","Legal rights, data-sharing boundaries, provenance, and governance."),
("Defensibility","Longitudinal context and operational history may become harder to replicate.","Real adoption, permissioned data, outcomes, and distribution.")
],[1.25,2.55,2.7])
callout("Strategic boundary", "The repository does not prove Talentbank partnerships, proprietary data access, customer demand, or the right to combine student, employer, and university records. Strategic fit is a hypothesis for leadership validation.")

h1("13. Pilot and Validation Plan")
h2("13.1 Pilot sequence")
for x in [
"Select one faculty, cohort, and employability owner.",
"Agree on two safe, actionable capability gaps.",
"Import a minimum controlled cohort dataset with explicit consent and manual verification.",
"Run one or two employer-backed interventions.",
"Capture student participation and evidence updates.",
"Review employer validation and warm-pipeline usefulness.",
"Produce one faculty action and evidence summary.",
"Conduct governance, fairness, and user-trust review before expansion."
]: numbered(x)
h2("13.2 Metrics")
table(["Area","Pilot metric"],[
("Activation","Eligible-student activation and Career Memory evidence completion."),
("Action","Intervention uptake, completion, and time from gap detection to action."),
("Employer","Partner participation, qualified challenge performers, and warm-candidate follow-up."),
("University","Staff time to produce an action/evidence summary and follow-through completion."),
("Trust","Consent comprehension, evidence correction requests, user confidence, and fairness concerns."),
("Retention","Return behaviour after deadlines, practice, rejection, intervention, or new evidence.")
],[1.35,5.15])
h2("13.3 Success criterion")
callout("Prove the workflow before the network", "The pilot should demonstrate that preserved evidence leads to a clearer next decision and completed action. It should not attempt to prove national-scale intelligence, predictive accuracy, or a complete marketplace.")

h1("14. Product Roadmap")
table(["Stage","Product focus","Exit criteria"],[
("Now — Round 2 prototype","Stabilise three workspaces, reliable demo paths, truth labels, and selected handoffs.","Consistent build, reliable flows, no unsupported claims, known fallback behaviour."),
("Pilot readiness","Consent, role access, provenance, auditability, controlled import, and pilot instrumentation.","Approved data model, governance review, pilot owner, baseline measures."),
("Bounded pilot","Faculty intervention loop with employer participation and student evidence.","Measured activation, follow-through, staff effort, trust, and workflow usefulness."),
("Post-pilot expansion","Deeper readiness, collaboration, employer pipeline, alumni, and accreditation workflows.","Evidence that each expansion solves a paid operational problem."),
("Future network","Permissioned integrations and longitudinal outcome learning.","Secure integration, validated models, clear governance, and demonstrated demand.")
],[1.25,3.15,2.1])
h2("14.1 Priority product improvements")
for x in [
"Add persistent demo-data and simulated-action labels to scores, sends, exports, and agent outputs.",
"Complete one challenge → Career Memory → employer profile → university outcome story.",
"Add evidence provenance and recency to employer-visible context.",
"Strengthen rejection-to-diagnosis and mentorship-to-Memory loops or keep them secondary.",
"Make current-cohort intervention explicit on Curriculum–Market Alignment.",
"Show accreditation reuse with source, period, owner, and validity.",
"Clarify source, prepared action, required approval, and external effect inside the AI Office."
]: bullet(x)

h1("15. Setup and Demo Guidance")
h2("15.1 Frontend")
for x in [
"Install dependencies: npm install",
"Start development: npm.cmd run dev",
"Production build check: npm.cmd run build",
"Open the local Vite URL and choose the intended workspace through the landing flow."
]: bullet(x)
h2("15.2 Backend")
para("The backend is partial and environment-dependent. Configure database and model credentials locally before running FastAPI services. Demo planning should define whether the frontend uses a live endpoint or its deterministic fallback and avoid switching state during a judged walkthrough.")
h2("15.3 Recommended stable demo")
for x in [
"Student: Career Memory → Career Intelligence → one preparation or opportunity action.",
"Employer: evidence-aware candidate detail → targeted validation → Campus Pipeline reactivation.",
"University: curriculum/readiness gap → current-cohort intervention → evidence pack / Accreditation Hub.",
"AI Office close: prepared action with explicit human approval and simulated-status language."
]: numbered(x)
h2("15.4 Avoid in the main demo")
para("Avoid placeholder settings/help routes, legacy aliases, unsupported exports, unconfigured backend calls, generic metrics without provenance, and any interaction that implies a live external send or production integration.")

h1("Appendix A. Route and Capability Inventory")
table(["Workspace","Current primary routes"],[
("Student","/student/home; /student/profile; /student/intelligence; /student/opportunities; /student/applications; /student/communities; /student/mentorships; /student/skill-development; /student/ai-companion"),
("Employer","/employer/home; /employer/talent; /employer/talent-discovery; /employer/posting; /employer/campus-pipeline; /employer/analytics; /employer/candidates; /employer/marketplace"),
("University","/university/overview; /university/student-readiness; /university/curriculum-alignment; /university/alumni-signals; /university/collaboration; /university/accreditation"),
("Legacy / placeholders","Several redirects and placeholder settings, reports, learning, and help routes remain in the repository and should not define the product narrative.")
],[1.1,5.4])
h2("Backend surface observed")
para("FastAPI includes user, candidate, and employer routers in backend/app/app.py. Additional router files and agent experiments are present, but not every file is included in the active application or connected to the Round 2 frontend.")

h1("Appendix B. Claim Language")
table(["Avoid","Use instead"],[
("Verified evidence","Evidence-oriented profile; the interface models verification and provenance."),
("AI predicts success","AI-supported recommendation; predictive accuracy is not validated."),
("Shared intelligence network","The architecture proposes a shared network; current workflows mostly use local or mock state."),
("Saves time / improves ROI","Designed to reduce repeated work; measure time and economic value in a pilot."),
("Autonomous AI employees","Specialised assistant interaction with human approval; autonomy is future-facing."),
("Automated accreditation","Accreditation readiness, evidence reuse, ownership, and draft preparation."),
("Talentbank advantage","Potential strategic fit pending validation of assets, rights, demand, and distribution."),
("Customers / partnerships / adoption","Do not claim unless documented evidence is provided.")
],[2.05,4.45])
h2("Source basis")
for x in [
"CareerOS Deep Value Discovery, docs/pitch/CAREEROS_DEEP_VALUE_DISCOVERY.md.",
"Round 1 submission, [Team GongBaiWan] CareerOS_Documentation_V1.pdf.",
"Refined editable documentation source, CareerOS_Product_Documentation_Refined.docx.",
"Current repository routes, components, stores, services, mock data, and backend files inspected in July 2026."
]: bullet(x)

doc.core_properties.title = "CareerOS Product Documentation — Round 2"
doc.core_properties.subject = "Product documentation grounded in the current CareerOS prototype"
doc.core_properties.author = "Team GongBaiWan"
doc.core_properties.keywords = "CareerOS, Talentbank, product documentation, Round 2"
doc.save(OUT)
print("CareerOS_Product_Documentation_Round_2.docx")
